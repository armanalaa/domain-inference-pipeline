#!/usr/bin/env python3
"""
extract_knowledge.py
────────────────────
Generate a business knowledge document (knowledge.docx) for the CCM /
TwoLevelDomainDiscovery pipeline.

Supports four LLM backends — select with --backend:

  ollama      Local Ollama (default, free, no API key)
              Requires: ollama serve  (in a separate terminal)

  anthropic   Anthropic Claude API
              Requires: --api_key sk-ant-...
              Install:  pip install anthropic

  openai      OpenAI API
              Requires: --api_key sk-...
              Install:  pip install openai

  groq        Groq API (free tier available)
              Requires: --api_key gsk_...  (free at console.groq.com)
              Install:  pip install groq

All model constraints (timeout, output tokens, context window, temperature,
tokens-per-minute rate limit, schema/source size limits) are exposed as CLI
arguments so you can tune them for any model without editing the script.

Usage
─────
    # Ollama — default settings
    python extract_knowledge.py --database eICU ^
        --pdf_dir knowledge/ --csv_dir csv/ --schema schema.json

    # Groq free tier — enforce 6000 TPM rate limit
    python extract_knowledge.py --database eICU --schema schema.json ^
        --backend groq --api_key gsk_... --tpm_limit 6000

    # Anthropic Claude — best quality
    python extract_knowledge.py --database eICU --schema schema.json ^
        --backend anthropic --api_key sk-ant-...

    # OpenAI with larger output cap
    python extract_knowledge.py --database eICU --schema schema.json ^
        --backend openai --model gpt-4o --api_key sk-... --max_tokens 3000

    # Ollama with extended timeout, larger context, more output
    python extract_knowledge.py --database eICU --schema schema.json ^
        --timeout 1800 --num_ctx 8192 --max_tokens 2048

Dependencies
────────────
    pip install requests pdfplumber python-docx
    pip install anthropic          # only for --backend anthropic
    pip install openai             # only for --backend openai
    pip install groq               # only for --backend groq
"""

import argparse
import csv
import json
import os
import shutil
import sys
import textwrap
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from path_utils import resolve_dataset_dir

# ── optional imports ──────────────────────────────────────────────────────────
try:
    import requests
    _REQUESTS_OK = True
except ImportError:
    _REQUESTS_OK = False

try:
    import pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from pipeline_utils import setup_log_file, clean_extract_knowledge
    _UTILS = True
except ImportError:
    _UTILS = False


# ─────────────────────────────────────────────────────────────────────────────
# LLMConfig  —  all model-specific parameters in one place
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class LLMConfig:
    """
    All tunable LLM parameters.  Every field has a sensible default and
    is overridable from the CLI — no editing of the script needed.
    """
    # ── Identification ────────────────────────────────────────────────────────
    backend:     str   = "ollama"
    model:       str   = ""
    api_key:     str   = ""

    # ── Generation constraints (apply to ALL backends) ────────────────────────
    max_tokens:  int   = 1024   # Maximum output tokens
    temperature: float = 0.2   # 0.0 = deterministic, 1.0 = creative

    # ── Ollama-specific ───────────────────────────────────────────────────────
    num_ctx:     int   = 4096   # Context window size (tokens)
    timeout:     int   = 1200   # HTTP timeout (seconds, 20 min default)
    ollama_url:  str   = "http://localhost:11434"

    # ── Rate limiting  (cloud backends) ──────────────────────────────────────
    # Set tpm_limit to your plan's tokens-per-minute cap.
    # The script tracks tokens generated in the current 60-second window and
    # sleeps automatically if the limit is about to be exceeded.
    # 0 = no rate limiting.
    # Examples:  Groq free tier = 6000   |   most paid plans = 0 (unlimited)
    tpm_limit:   int   = 0

    # ── Prompt size limits ────────────────────────────────────────────────────
    schema_max_chars:  int = 1500   # Max chars of schema summary in the prompt
    source_max_chars:  int = 2000   # Max chars of PDF/CSV content per LLM call

    # ── Chunked PDF processing ────────────────────────────────────────────────
    # When pdf_chunk_size > 0, PDFs are split into chunks of this many chars
    # and the LLM is called once per chunk. All sections are merged into one
    # knowledge.docx. Set 0 to disable (send everything in one call).
    # Groq free tier safe value: ~20000 chars per chunk.
    pdf_chunk_size: int = 0


# ─────────────────────────────────────────────────────────────────────────────
# Backend defaults
# ─────────────────────────────────────────────────────────────────────────────

BACKEND_DEFAULT_MODELS = {
    "ollama":    None,
    "anthropic": "claude-sonnet-4-6",
    "openai":    "gpt-4o-mini",
    "groq":      "llama-3.3-70b-versatile",
}

OLLAMA_CANDIDATE_MODELS = [
    "mistral-ctx4k",    # custom Modelfile — preferred
    "mistral:latest",
    "mistral",
    "llama3:latest",
    "llama3",
    "llama2:latest",
    "llama2",
]


# ─────────────────────────────────────────────────────────────────────────────
# PROGRESS BAR
# ─────────────────────────────────────────────────────────────────────────────

class TokenProgressBar:
    """
    Single-line progress bar driven by token count.
    Total = max_tokens cap.  Saturates at 99% until finish() is called.
    Shows: % | bar | tokens | elapsed | ETA | current section heading
    """
    def __init__(self, total: int):
        self.total    = max(total, 1)
        self.current  = 0
        self.start_t  = time.time()
        self._section = ""
        self._width   = shutil.get_terminal_size(fallback=(80, 24)).columns
        self._draw()

    def update(self, n: int = 1, section: str = ""):
        self.current += n
        if section:
            self._section = section
        self._draw()

    def finish(self):
        self.current = self.total
        self._draw()
        sys.stdout.write("\n")
        sys.stdout.flush()

    def _draw(self):
        elapsed = time.time() - self.start_t
        pct     = min(self.current / self.total, 0.99)
        eta     = (_fmt_time(elapsed / pct * (1 - pct))
                   if self.current > 0 and pct < 0.99 else "--:--")
        prefix  = f"{int(pct * 100):3d}%"
        suffix  = (f" {self.current}/{self.total}tok"
                   f" {_fmt_time(elapsed)}<{eta}"
                   + (f"  {self._section}" if self._section else ""))
        bar_w   = max(self._width - len(prefix) - len(suffix) - 4, 4)
        filled  = int(bar_w * pct)
        line    = f"\r{prefix} |{'█'*filled + '░'*(bar_w-filled)}| {suffix}"
        sys.stdout.write(line[:self._width - 1])
        sys.stdout.flush()


def _fmt_time(s: float) -> str:
    s = int(s); m, s = divmod(s, 60); h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _detect_section(chunks: list) -> str:
    joined = "".join(chunks[-30:])
    idx    = joined.rfind("## ")
    if idx != -1:
        c = joined[idx + 3:].split("\n")[0].strip()
        return c[:38] if c else ""
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# RATE LIMITER
# ─────────────────────────────────────────────────────────────────────────────

class RateLimiter:
    """
    Token-per-minute rate limiter.
    Tracks tokens in the current 60-second window.  Sleeps when the cap is
    reached.  tpm_limit=0 disables limiting entirely.
    """
    def __init__(self, tpm_limit: int):
        self.tpm_limit   = tpm_limit
        self._win_start  = time.time()
        self._win_tokens = 0

    def consume(self, n: int = 1):
        if self.tpm_limit <= 0:
            return
        now = time.time()
        if now - self._win_start >= 60.0:
            self._win_start  = now
            self._win_tokens = 0
        self._win_tokens += n
        if self._win_tokens >= self.tpm_limit:
            sleep_s = 60.0 - (now - self._win_start) + 1.0
            if sleep_s > 0:
                sys.stdout.write(
                    f"\n[rate]  TPM limit ({self.tpm_limit}) reached — "
                    f"sleeping {sleep_s:.0f}s ...\n")
                sys.stdout.flush()
                time.sleep(sleep_s)
                self._win_start  = time.time()
                self._win_tokens = 0


# ─────────────────────────────────────────────────────────────────────────────
# MODEL RESOLUTION
# ─────────────────────────────────────────────────────────────────────────────

def _check_ollama(url: str) -> bool:
    try:
        requests.get(url, timeout=5); return True
    except Exception:
        return False


def _list_ollama_models(url: str) -> set:
    try:
        r = requests.get(url.rstrip("/") + "/api/tags", timeout=10)
        if r.status_code == 200:
            return {m["name"] for m in r.json().get("models", [])}
    except Exception:
        pass
    return set()


def resolve_model(cfg: LLMConfig) -> str:
    if cfg.backend == "ollama":
        if not _check_ollama(cfg.ollama_url):
            print("[ERROR] Ollama is not running.")
            print("        Open a NEW terminal and run:  ollama serve")
            sys.exit(1)
        print(f"[ollama] Connected at {cfg.ollama_url}")
        if cfg.model:
            print(f"[model]  Using: {cfg.model}"); return cfg.model
        available = _list_ollama_models(cfg.ollama_url)
        print(f"[model]  Available locally: {', '.join(sorted(available)) or 'none'}")
        for m in OLLAMA_CANDIDATE_MODELS:
            if m in available:
                print(f"[model]  Selected: {m}"); return m
        print("[ERROR] No suitable Ollama model found.")
        print("        Run:  ollama pull mistral")
        print("        Then: ollama create mistral-ctx4k -f Modelfile")
        sys.exit(1)
    model = cfg.model or BACKEND_DEFAULT_MODELS[cfg.backend]
    print(f"[model]  Backend={cfg.backend}  Model={model}")
    return model


def _resolve_api_key(backend: str, explicit: str) -> str:
    if explicit:
        return explicit
    env = {"anthropic": "ANTHROPIC_API_KEY",
           "openai":    "OPENAI_API_KEY",
           "groq":      "GROQ_API_KEY"}.get(backend, "")
    key = os.environ.get(env, "")
    if key: print(f"[auth]   Using API key from {env}")
    return key


# ─────────────────────────────────────────────────────────────────────────────
# SOURCE MATERIAL READERS
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# SOURCE MATERIAL READERS
# ─────────────────────────────────────────────────────────────────────────────

def read_pdfs_chunked(pdf_dir: Optional[Path],
                      chars_per_chunk: int) -> list[str]:
    """
    Read all PDFs in pdf_dir and split into chunks of ~chars_per_chunk chars.
    Returns a list of text chunks — each chunk fits within the LLM context.
    Pages are never split mid-page; chunks accumulate whole pages until the
    size limit is reached, then a new chunk starts.
    """
    if not pdf_dir or not pdf_dir.exists(): return []
    if not _PDF_OK:
        print("[source] pdfplumber not installed — Run: pip install pdfplumber")
        return []

    all_pages = []   # list of (filename, page_num, text)
    for path in sorted(pdf_dir.glob("*.pdf")):
        try:
            with pdfplumber.open(path) as pdf:
                n = len(pdf.pages)
                print(f"  [pdf] {path.name}: {n} pages", end="", flush=True)
                extracted = 0
                for i, page in enumerate(pdf.pages):
                    text = (page.extract_text() or "").strip()
                    if text:
                        all_pages.append((path.name, i + 1, text))
                        extracted += 1
                print(f" ({extracted} non-empty)")
        except Exception as exc:
            print(f"\n  [pdf] {path.name}: ERROR — {exc}")

    if not all_pages:
        return []

    # Group pages into chunks
    chunks   = []
    current  = []
    cur_size = 0

    for fname, pnum, text in all_pages:
        page_text = f"[PDF: {fname} p.{pnum}]\n{text}"
        if cur_size + len(page_text) > chars_per_chunk and current:
            chunks.append("\n\n".join(current))
            current  = []
            cur_size = 0
        current.append(page_text)
        cur_size += len(page_text)

    if current:
        chunks.append("\n\n".join(current))

    total_chars = sum(len(c) for c in chunks)
    print(f"  [pdf] {len(all_pages)} pages → {len(chunks)} chunks "
          f"({total_chars:,} chars total)")
    return chunks


def read_pdfs(pdf_dir: Optional[Path]) -> str:
    """Read all PDFs as a single string (legacy, used when chunking is off)."""
    chunks = read_pdfs_chunked(pdf_dir, chars_per_chunk=10_000_000)
    return "\n\n".join(chunks)


def read_texts(pdf_dir: Optional[Path]) -> str:
    """Read all .txt files from the same source folder as PDFs."""
    if not pdf_dir or not pdf_dir.exists(): return ""
    chunks = []
    for path in sorted(pdf_dir.glob("*.txt")):
        try:
            text = path.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                chunks.append(f"[TXT: {path.name}]\n{text}")
                print(f"  [txt] {path.name}: {len(text):,} chars")
        except Exception as exc:
            print(f"  [txt] {path.name}: ERROR — {exc}")
    return "\n\n".join(chunks)


def read_csvs(csv_dir: Optional[Path], max_rows: int = 3) -> str:
    if not csv_dir or not csv_dir.exists(): return ""
    chunks = []
    for path in sorted(csv_dir.glob("*.csv")):
        try:
            with open(path, newline="", encoding="utf-8-sig") as fh:
                rows = []
                for i, row in enumerate(csv.reader(fh)):
                    if i > max_rows: break
                    rows.append(", ".join(str(c) for c in row))
            if rows:
                chunks.append(f"[CSV: {path.name}]\n" + "\n".join(rows))
                print(f"  [csv] {path.name}: {len(rows)} rows shown")
        except Exception as exc:
            print(f"  [csv] {path.name}: ERROR — {exc}")
    return "\n\n".join(chunks)


def build_schema_summary(schema_path: Optional[Path]) -> str:
    if not schema_path or not schema_path.exists(): return ""
    try:
        with open(schema_path, encoding="utf-8") as fh:
            schema = json.load(fh)
    except Exception as exc:
        print(f"[schema] Could not read {schema_path}: {exc}"); return ""
    lines = []
    for tname, tbody in schema.get("tables", {}).items():
        cols = [c["name"] for c in tbody.get("columns", [])]
        lines.append(f"{tname} ({', '.join(cols)})")
    summary = "\n".join(lines)
    print(f"[schema] {len(schema.get('tables',{}))} tables, "
          f"{len(summary):,} chars before truncation")
    return summary


# ─────────────────────────────────────────────────────────────────────────────
# PROMPTS
# ─────────────────────────────────────────────────────────────────────────────

SYSTEM_MESSAGE = textwrap.dedent("""\
    You are a senior business analyst with deep expertise in enterprise process
    modelling across multiple industries. You write structured knowledge documents
    that describe HOW a business operates — the activities people perform, the
    decisions they make, and the information they produce or consume.

    Your output is used as input to an automated domain-discovery pipeline
    (CCM / TwoLevelDomainDiscovery). The pipeline extracts business concepts
    from your text and uses them to group related database columns into semantic
    domains. Your goal is to write richly descriptive BUSINESS PROCESS text so
    that thematically related columns cluster into coherent business domains.

    Rules you must follow without exception:
    1. Write about what PEOPLE DO, not what TABLES STORE.
       Bad:  "The Product table stores product information including ProductID."
       Good: "Product managers register new items in the catalog by recording
              the product code, name, category, and list price."
    2. Never start a bullet with "The [Table] table stores/records/contains".
       Always start with an organisational role or a business action verb.
       Good starts: "Sales representatives ...", "HR managers ...",
                    "Customers place ...", "Warehouse staff track ...",
                    "Purchasing teams issue ...", "Finance records ..."
    3. Mention specific column names or table names only as supporting detail
       at the END of a bullet, not as the subject of the sentence.
    4. Use plain business vocabulary — no technical database jargon.
    5. Name the organisational roles that perform each activity in Performed By.
    6. Write exactly 3 to 5 activity bullets per section (one sentence each).
    7. Use the section format shown below — no deviation.
    8. Do NOT output anything except the formatted knowledge document.\
""")

USER_TEMPLATE = textwrap.dedent("""\
    Generate a structured business workflow document for the database described
    below. Focus on what BUSINESS USERS DO, not on what database tables contain.

    DATABASE NAME: {database_name}

    SCHEMA SUMMARY (use as reference for column/table names only):
    {schema_summary}

    SOURCE MATERIALS (use to understand the business domain):
    {source_material_text}

    ---

    Output format — use this EXACT structure for every section:

    # {database_name} — Business Workflows

    ## 1. [Business Process Name]

    Performed By: [Role1], [Role2]

    Activity:
    - [Business action sentence starting with a role or verb, ending with column/table reference.]
    - [...]
    (3-5 bullets per section)

    ## 2. [Next Business Process Name]
    ...

    ---

    Instructions:
    - Identify 7 to 10 distinct business processes (e.g. Order Placement,
      Employee Onboarding, Inventory Replenishment, Invoice Processing).
    - Each bullet must describe a human activity or decision — never a database
      operation. Do NOT write "Table X stores Y". Write "Role Z performs Y."
    - Order sections from core entity management through transactional workflows
      to support and analytical processes.
    - Keep each bullet to ONE sentence. Be concise — total output must fit in
      1000 tokens.
    - Do not number bullets; use a hyphen (-) only.
    - Do not add preamble, summary, or closing remarks.
    - Output plain text only — no Markdown beyond ##, no bold, no tables.\
""")


def _build_messages(database_name: str, schema_summary: str,
                    source_text: str, cfg: LLMConfig):
    schema_t = schema_summary[:cfg.schema_max_chars]
    source_t = source_text[:cfg.source_max_chars]
    user_msg = USER_TEMPLATE.format(
        database_name        = database_name,
        schema_summary       = schema_t or "(no schema.json provided)",
        source_material_text = source_t or "(no source material provided)",
    )
    def _trunc_note(label, orig, trunc):
        return (f"      {label}: {len(trunc):,} chars (truncated from {len(orig):,})"
                if len(orig) > len(trunc)
                else f"      {label}: {len(trunc):,} chars")
    print(_trunc_note("schema_summary ", schema_summary, schema_t))
    print(_trunc_note("source_material", source_text,    source_t))
    total = len(SYSTEM_MESSAGE) + len(user_msg)
    print(f"      total prompt    : {total:,} chars (~{total//3:,} tokens)")
    return SYSTEM_MESSAGE, user_msg


# ─────────────────────────────────────────────────────────────────────────────
# LLM BACKENDS
# ─────────────────────────────────────────────────────────────────────────────

def _open_inc(path: Optional[Path]):
    if path:
        path.parent.mkdir(parents=True, exist_ok=True)
        return open(path, "w", encoding="utf-8")
    return None


def _tok(token: str, chunks: list, fh, bar: TokenProgressBar,
         limiter: RateLimiter):
    """Append token to chunks, write to file, update bar, apply rate limit."""
    chunks.append(token)
    if fh:
        fh.write(token); fh.flush()
    bar.update(1, section=_detect_section(chunks))
    limiter.consume(1)


def _print_cfg(cfg: LLMConfig, inc_path: Optional[Path]):
    print(f"      max_tokens      : {cfg.max_tokens}")
    print(f"      temperature     : {cfg.temperature}")
    print(f"      tpm_limit       : "
          f"{cfg.tpm_limit:,} tok/min" if cfg.tpm_limit > 0 else "      tpm_limit       : none")
    if cfg.backend == "ollama":
        print(f"      num_ctx         : {cfg.num_ctx}")
        print(f"      timeout         : {cfg.timeout}s ({cfg.timeout//60} min)")
    if inc_path:
        print(f"      live output     : {inc_path}")
    print()


def _call_ollama(cfg: LLMConfig, sys_msg: str, usr_msg: str,
                 inc_path: Optional[Path]) -> str:
    gen_url = cfg.ollama_url.rstrip("/") + "/api/generate"
    payload = {
        "model":   cfg.model,
        "prompt":  sys_msg + "\n\n" + usr_msg,
        "stream":  True,
        "options": {
            "temperature": cfg.temperature,
            "num_predict": cfg.max_tokens,
            "num_ctx":     cfg.num_ctx,
        },
    }
    bar     = TokenProgressBar(cfg.max_tokens)
    limiter = RateLimiter(cfg.tpm_limit)
    chunks  = []; fh = _open_inc(inc_path)
    try:
        with requests.post(gen_url, json=payload,
                           stream=True, timeout=cfg.timeout) as resp:
            if resp.status_code != 200:
                bar.finish()
                print(f"\n[ERROR] Ollama {resp.status_code}: {resp.text[:300]}")
                sys.exit(1)
            for raw in resp.iter_lines():
                if not raw: continue
                try: data = json.loads(raw)
                except json.JSONDecodeError: continue
                if t := data.get("response", ""):
                    _tok(t, chunks, fh, bar, limiter)
                if data.get("done"): break
    except requests.exceptions.ConnectionError:
        bar.finish()
        print("\n[ERROR] Lost connection to Ollama. Is 'ollama serve' running?")
        sys.exit(1)
    except requests.exceptions.Timeout:
        bar.finish()
        print(f"\n[ERROR] Timed out after {cfg.timeout}s.")
        print(f"        Try: --timeout {cfg.timeout + 600}")
        sys.exit(1)
    finally:
        if fh: fh.close()
    bar.finish()
    return "".join(chunks)


def _call_anthropic(cfg: LLMConfig, sys_msg: str, usr_msg: str,
                    inc_path: Optional[Path]) -> str:
    try: import anthropic as _ant
    except ImportError:
        print("[ERROR] Run: pip install anthropic"); sys.exit(1)
    bar = TokenProgressBar(cfg.max_tokens)
    limiter = RateLimiter(cfg.tpm_limit)
    chunks = []; fh = _open_inc(inc_path)
    try:
        with _ant.Anthropic(api_key=cfg.api_key).messages.stream(
            model=cfg.model, max_tokens=cfg.max_tokens,
            temperature=cfg.temperature, system=sys_msg,
            messages=[{"role": "user", "content": usr_msg}],
        ) as stream:
            for t in stream.text_stream:
                _tok(t, chunks, fh, bar, limiter)
    except Exception as exc:
        bar.finish(); print(f"\n[ERROR] Anthropic: {exc}"); sys.exit(1)
    finally:
        if fh: fh.close()
    bar.finish()
    return "".join(chunks)


def _call_openai(cfg: LLMConfig, sys_msg: str, usr_msg: str,
                 inc_path: Optional[Path]) -> str:
    try: import openai as _oai
    except ImportError:
        print("[ERROR] Run: pip install openai"); sys.exit(1)
    bar = TokenProgressBar(cfg.max_tokens)
    limiter = RateLimiter(cfg.tpm_limit)
    chunks = []; fh = _open_inc(inc_path)
    try:
        stream = _oai.OpenAI(api_key=cfg.api_key).chat.completions.create(
            model=cfg.model, stream=True, max_tokens=cfg.max_tokens,
            temperature=cfg.temperature,
            messages=[{"role": "system", "content": sys_msg},
                      {"role": "user",   "content": usr_msg}],
        )
        for chunk in stream:
            if t := chunk.choices[0].delta.content or "":
                _tok(t, chunks, fh, bar, limiter)
    except Exception as exc:
        bar.finish(); print(f"\n[ERROR] OpenAI: {exc}"); sys.exit(1)
    finally:
        if fh: fh.close()
    bar.finish()
    return "".join(chunks)


def _call_groq(cfg: LLMConfig, sys_msg: str, usr_msg: str,
               inc_path: Optional[Path]) -> str:
    try: from groq import Groq as _Groq
    except ImportError:
        print("[ERROR] Run: pip install groq"); sys.exit(1)
    bar = TokenProgressBar(cfg.max_tokens)
    limiter = RateLimiter(cfg.tpm_limit)
    chunks = []; fh = _open_inc(inc_path)
    try:
        stream = _Groq(api_key=cfg.api_key).chat.completions.create(
            model=cfg.model, stream=True, max_tokens=cfg.max_tokens,
            temperature=cfg.temperature,
            messages=[{"role": "system", "content": sys_msg},
                      {"role": "user",   "content": usr_msg}],
        )
        for chunk in stream:
            if t := chunk.choices[0].delta.content or "":
                _tok(t, chunks, fh, bar, limiter)
    except Exception as exc:
        bar.finish(); print(f"\n[ERROR] Groq: {exc}"); sys.exit(1)
    finally:
        if fh: fh.close()
    bar.finish()
    return "".join(chunks)


# ─────────────────────────────────────────────────────────────────────────────
# UNIFIED DISPATCHER
# ─────────────────────────────────────────────────────────────────────────────

def call_llm(cfg: LLMConfig, database_name: str,
             schema_summary: str, source_text: str,
             incremental_path: Optional[Path] = None) -> str:

    sys_msg, usr_msg = _build_messages(database_name, schema_summary,
                                       source_text, cfg)
    print(f"\n[llm] Backend: {cfg.backend}  |  Model: {cfg.model}")
    _print_cfg(cfg, incremental_path)

    t0 = time.time()
    fn = {"ollama": _call_ollama, "anthropic": _call_anthropic,
          "openai": _call_openai, "groq": _call_groq}.get(cfg.backend)
    if fn is None:
        print(f"[ERROR] Unknown backend: {cfg.backend}"); sys.exit(1)

    result  = fn(cfg, sys_msg, usr_msg, incremental_path)
    elapsed = time.time() - t0
    print(f"[llm] Done — {len(result):,} chars  |  "
          f"{elapsed:.0f}s ({elapsed/60:.1f} min)")

    if not result.strip():
        print("[ERROR] LLM returned an empty response."); sys.exit(1)
    return result


# ─────────────────────────────────────────────────────────────────────────────
# PARSE LLM OUTPUT
# ─────────────────────────────────────────────────────────────────────────────

def parse_sections(raw_text: str):
    lines, doc_title, sections, current = raw_text.splitlines(), "", [], None
    for line in lines:
        s = line.strip()
        if s.startswith("# ") and not s.startswith("## "):
            doc_title = s[2:].strip(); continue
        if s.startswith("## "):
            if current: sections.append(current)
            heading = s[3:].strip(); parts = heading.split(".", 1)
            try:    num, title = int(parts[0].strip()), (parts[1].strip() if len(parts)>1 else heading)
            except: num, title = len(sections)+1, heading
            current = {"number": num, "title": title, "performed_by": "", "bullets": []}
            continue
        if current is None: continue
        if s.lower().startswith("performed by:"):
            current["performed_by"] = s.split(":", 1)[1].strip(); continue
        if s.lower() == "activity:": continue
        if s.startswith("- "):
            b = s[2:].strip()
            if b: current["bullets"].append(b)
    if current: sections.append(current)
    for sec in sections:
        if not sec["bullets"]:
            print(f"  [parse] WARNING: section {sec['number']} '{sec['title']}' has no bullets")
    return doc_title, sections


# ─────────────────────────────────────────────────────────────────────────────
# WRITE knowledge.docx
# ─────────────────────────────────────────────────────────────────────────────

def _font(run, size_pt, bold=False, italic=False, color_rgb=None):
    run.font.name = "Calibri"; run.font.size = Pt(size_pt)
    run.font.bold = bold; run.font.italic = italic
    if color_rgb: run.font.color.rgb = RGBColor(*color_rgb)


def write_docx(doc_title, sections, database_name, output_path, raw_text):
    if not _DOCX_OK:
        txt = output_path.with_suffix(".txt")
        txt.write_text(raw_text, encoding="utf-8")
        print(f"[docx] python-docx not installed; plain text saved to {txt}")
        return

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    p = doc.add_paragraph()
    _font(p.add_run(doc_title or f"{database_name} — Business Workflows"),
          22, bold=True, color_rgb=(27, 58, 92))
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    _font(p.add_run(f"Knowledge Document for CCM / TwoLevelDomainDiscovery  |  "
                    f"Database: {database_name}"),
          10, italic=True, color_rgb=(85, 85, 85))
    p.paragraph_format.space_after = Pt(16)

    for sec in sections:
        p = doc.add_paragraph()
        _font(p.add_run(f"{sec['number']}. {sec['title']}"),
              13, bold=True, color_rgb=(27, 58, 92))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)

        if sec["performed_by"]:
            p = doc.add_paragraph()
            _font(p.add_run("Performed By: "), 10, bold=True)
            _font(p.add_run(sec["performed_by"]), 10)
            p.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        _font(p.add_run("Activity:"), 10, bold=True)
        p.paragraph_format.space_after = Pt(2)

        bullets = sec["bullets"] or ["[No bullets — LLM output may have been truncated]"]
        is_placeholder = not sec["bullets"]
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            _font(p.add_run(b), 10,
                  italic=is_placeholder,
                  color_rgb=((180, 0, 0) if is_placeholder else None))
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    p = doc.add_paragraph()
    _font(p.add_run(f"{database_name}  |  CCM Pipeline — TwoLevelDomainDiscovery  |  "
                    f"Auto-generated by extract_knowledge.py"),
          8, italic=True, color_rgb=(120, 120, 120))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    print(f"[docx] Written: {output_path}  ({output_path.stat().st_size/1024:.1f} KB)")
    try:
        chars = len("\n".join(p.text for p in Document(str(output_path)).paragraphs
                              if p.text.strip()))
        print(f"[docx] Pipeline will read: {chars:,} chars")
        if chars > 4000:
            print(f"[docx] WARNING: {chars:,} chars — Script 1 may truncate. "
                  f"Try reducing --max_tokens or bullet count.")
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def build_parser() -> argparse.ArgumentParser:
    D = LLMConfig()   # defaults
    p = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    g = p.add_argument_group("Source material")
    g.add_argument("--database", "-d", required=True,
                   help="Database / project name (e.g. eICU, Northwind).")
    g.add_argument("--pdf_dir", "-p", type=Path, default=None,
                   help="Folder with PDF documents. Optional.")
    g.add_argument("--csv_dir", "-c", type=Path, default=None,
                   help="Folder with CSV files (header + 3 rows). Optional.")
    g.add_argument("--schema",  "-s", type=Path, default=None,
                   help="Path to schema.json. Strongly recommended.")

    g = p.add_argument_group("LLM backend")
    g.add_argument("--backend", "-b", default="ollama",
                   choices=["ollama", "anthropic", "openai", "groq"],
                   help="LLM backend (default: ollama).")
    g.add_argument("--model", "-m", default=None,
                   help=(f"Model name. Defaults: ollama=auto, "
                         f"anthropic={BACKEND_DEFAULT_MODELS['anthropic']}, "
                         f"openai={BACKEND_DEFAULT_MODELS['openai']}, "
                         f"groq={BACKEND_DEFAULT_MODELS['groq']}"))
    g.add_argument("--api_key", "-k", default=None,
                   help="API key for cloud backends. "
                        "Also via ANTHROPIC_API_KEY / OPENAI_API_KEY / GROQ_API_KEY.")

    g = p.add_argument_group(
        "Model constraints  (tune per model/plan without editing the script)")
    g.add_argument("--max_tokens", type=int, default=D.max_tokens,
                   help=f"Max output tokens (default: {D.max_tokens}). "
                        "Increase for longer docs; decrease for speed/cost.")
    g.add_argument("--temperature", type=float, default=D.temperature,
                   help=f"Sampling temperature (default: {D.temperature}). "
                        "0=deterministic, 1=creative.")
    g.add_argument("--num_ctx", type=int, default=D.num_ctx,
                   help=f"Ollama context window tokens (default: {D.num_ctx}). "
                        "Ignored for cloud backends.")
    g.add_argument("--timeout", type=int, default=D.timeout,
                   help=f"Ollama HTTP timeout seconds (default: {D.timeout}). "
                        "Increase for slow machines.")
    g.add_argument("--tpm_limit", type=int, default=D.tpm_limit,
                   help="Tokens-per-minute rate limit (default: 0 = none). "
                        "Set to your plan cap to avoid 429 errors. "
                        "Example: Groq free = 6000.")
    g.add_argument("--schema_max_chars", type=int, default=D.schema_max_chars,
                   help=f"Max schema summary chars in prompt (default: {D.schema_max_chars}).")
    g.add_argument("--source_max_chars", type=int, default=D.source_max_chars,
                   help=f"Max PDF/CSV content chars in prompt (default: {D.source_max_chars}).")
    g.add_argument("--pdf_chunk_size", type=int, default=D.pdf_chunk_size,
                   help="Split PDFs into chunks of this many chars and call the LLM "
                        "once per chunk — merges all sections into one knowledge.docx. "
                        "0 = disabled (send all at once). "
                        "Recommended for large PDFs with Groq free tier: 20000.")

    g = p.add_argument_group("Output / Ollama URL")
    g.add_argument("--output", "-o", type=Path, default=Path("knowledge.docx"),
                   help="Output .docx path (default: knowledge.docx).")
    g.add_argument("--ollama_url", default=None,
                   help=f"Ollama base URL (default: {D.ollama_url}).")
    g.add_argument("--dataset_dir", default=None,
                   help="Dataset working directory (schema.json, knowledge/, csv/). "
                        "Use when extract_knowledge.py lives in a parent folder: "
                        "--dataset_dir Chinook. Defaults to current working directory.")
    g.add_argument("--log_dir", default="logs",
                   help="Folder for log files (default: ./logs).")
    g.add_argument("--no_clean", action="store_true",
                   help="Skip removal of knowledge.docx and chunks/ before running.")
    return p


def main() -> int:
    parser = build_parser()
    args   = parser.parse_args()

    # ── Dataset directory — chdir so all relative paths resolve correctly ────
    if args.dataset_dir is not None:
        import os as _os
        _os.chdir(resolve_dataset_dir(args.dataset_dir))

    # ── Logging ───────────────────────────────────────────────────────────────
    if _UTILS:
        setup_log_file("extract_knowledge", log_dir=args.log_dir)

    # ── Cleanup ───────────────────────────────────────────────────────────────
    if not args.no_clean:
        if _UTILS:
            clean_extract_knowledge(
                output  = args.output,
                pdf_dir = args.pdf_dir,
            )
        else:
            # Fallback without pipeline_utils
            for p in [args.output,
                      args.output.with_suffix(".txt")]:
                if Path(p).exists():
                    Path(p).unlink()
                    print(f"[clean] Removed {p}")
            if args.pdf_dir:
                chunks = Path(args.pdf_dir) / "chunks"
                if chunks.exists():
                    import shutil as _sh
                    _sh.rmtree(chunks)
                    print(f"[clean] Removed {chunks}")

    if not _REQUESTS_OK:
        print("ERROR: requests not installed.  Run: pip install requests")
        return 1

    cfg = LLMConfig(
        backend          = args.backend,
        model            = args.model or "",
        api_key          = _resolve_api_key(args.backend, args.api_key),
        max_tokens       = args.max_tokens,
        temperature      = args.temperature,
        num_ctx          = args.num_ctx,
        timeout          = args.timeout,
        ollama_url       = (args.ollama_url or LLMConfig().ollama_url).rstrip("/"),
        tpm_limit        = args.tpm_limit,
        schema_max_chars = args.schema_max_chars,
        source_max_chars = args.source_max_chars,
        pdf_chunk_size   = args.pdf_chunk_size,
    )

    if cfg.backend != "ollama" and not cfg.api_key:
        env = {"anthropic": "ANTHROPIC_API_KEY",
               "openai":    "OPENAI_API_KEY",
               "groq":      "GROQ_API_KEY"}[cfg.backend]
        print(f"[ERROR] --backend {cfg.backend} requires an API key.")
        print(f"        Pass --api_key <key>  or  set {env}=<key>")
        return 1

    cfg.model = resolve_model(cfg)

    print("\n[source] Reading source materials ...")
    txt_text       = read_texts(args.pdf_dir)
    csv_text       = read_csvs(args.csv_dir)
    schema_summary = build_schema_summary(args.schema)
    static_source  = "\n\n".join(filter(None, [txt_text, csv_text]))

    if not args.pdf_dir and not static_source and not schema_summary:
        print("WARNING: No source material — output quality will be low.")

    # ── Chunked or single-call mode ───────────────────────────────────────────
    raw_text = ""   # will be set in single-call mode; chunked mode uses sections directly
    if cfg.pdf_chunk_size > 0 and args.pdf_dir and args.pdf_dir.exists():
        # Chunked mode: split PDFs into pages, feed each chunk to LLM separately
        pdf_chunks = read_pdfs_chunked(args.pdf_dir, cfg.pdf_chunk_size)

        if not pdf_chunks:
            # No PDFs — fall back to single call with static sources
            pdf_chunks = [""]

        all_sections = []
        seen_titles  = set()
        doc_title    = ""

        total_chunks = len(pdf_chunks)
        print(f"\n[chunk] Processing {total_chunks} chunk(s) — "
              f"one LLM call per chunk ...")

        # Create chunks subfolder inside pdf_dir
        chunks_dir = args.pdf_dir / "chunks"
        chunks_dir.mkdir(parents=True, exist_ok=True)
        print(f"[chunk] Chunk files will be saved to: {chunks_dir}")

        for idx, chunk in enumerate(pdf_chunks, 1):
            print(f"\n[chunk] ── Chunk {idx}/{total_chunks} "
                  f"({len(chunk):,} chars) ──────────────────")

            source_for_call = "\n\n".join(filter(None, [chunk, static_source]))

            try:
                raw = call_llm(
                    cfg              = cfg,
                    database_name    = args.database,
                    schema_summary   = schema_summary,
                    source_text      = source_for_call,
                    incremental_path = chunks_dir / f"chunk{idx:02d}.txt",
                )
            except SystemExit as e:
                print(f"\n[chunk] WARNING: Chunk {idx} failed (exit code {e.code}) — skipping.")
                print(f"[chunk] Tip: reduce --pdf_chunk_size or --source_max_chars and retry.")
                continue

            title, secs = parse_sections(raw)
            if title and not doc_title:
                doc_title = title

            # Deduplicate sections by title — keep first occurrence
            new = 0
            for sec in secs:
                key = sec["title"].lower().strip()
                if key not in seen_titles and sec["bullets"]:
                    seen_titles.add(key)
                    # Renumber sequentially
                    sec["number"] = len(all_sections) + 1
                    all_sections.append(sec)
                    new += 1
            print(f"[chunk] {len(secs)} sections found, {new} new added "
                  f"(total so far: {len(all_sections)})")

        # Use the merged sections
        sections = all_sections
        print(f"\n[merge] Final merged document: {len(sections)} sections")

    else:
        # Single-call mode (original behaviour)
        pdf_text    = read_pdfs(args.pdf_dir)
        source_text = "\n\n".join(filter(None, [pdf_text, static_source]))

        raw_text = call_llm(
            cfg              = cfg,
            database_name    = args.database,
            schema_summary   = schema_summary,
            source_text      = source_text,
            incremental_path = Path(args.database.lower().replace(" ", "_") + "_knowledge.txt"),
        )
        doc_title, sections = parse_sections(raw_text)
    print(f"\n[parse] Found {len(sections)} sections:")
    for sec in sections:
        by = f", Performed By: {sec['performed_by']}" if sec["performed_by"] else ""
        print(f"         {sec['number']:2d}. {sec['title']}  "
              f"({len(sec['bullets'])} bullets{by})")

    if len(sections) < 5:
        print(f"\n[parse] WARNING: only {len(sections)} sections (expected 7-10).")
        if cfg.backend == "ollama":
            print(f"         Try: --timeout {cfg.timeout+600} --max_tokens {cfg.max_tokens+512}")
        else:
            print("         Try: --max_tokens 2048 or a larger --model")

    # Build raw_text fallback from sections (used if python-docx not installed)
    if not raw_text:
        raw_text = "\n\n".join(
            f"## {s['number']}. {s['title']}\nPerformed By: {s['performed_by']}\nActivity:\n" +
            "\n".join(f"- {b}" for b in s["bullets"])
            for s in sections
        )

    write_docx(doc_title=doc_title, sections=sections, database_name=args.database,
               output_path=args.output, raw_text=raw_text)

    print("\n[done] knowledge.docx is ready for the CCM pipeline.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
